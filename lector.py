import docx
from deep_translator import GoogleTranslator

# 1. Leemos el documento original en español
documento = docx.Document('prueba.docx')
texto_extraido = []

for parrafo in documento.paragraphs:
    if parrafo.text.strip(): 
        texto_extraido.append(parrafo.text)

texto_completo = '\n'.join(texto_extraido)

print("--- TEXTO ORIGINAL ---")
print(texto_completo)
print("-" * 30)

# 2. Traducimos el texto automáticamente a Inglés
print("Traduciendo documento... Por favor espera...")
texto_traducido = GoogleTranslator(source='es', target='en').translate(texto_completo)

print("\n--- TEXTO TRADUCIDO ---")
print(texto_traducido)
print("-" * 30)

# 3. Creamos un nuevo documento de Word y guardamos la traducción
print("Guardando traducción en un nuevo archivo...")
documento_traducido = docx.Document()

# Separamos el texto por líneas para que respete los párrafos originales
for linea in texto_traducido.split('\n'):
    documento_traducido.add_paragraph(linea)

# Guardamos el archivo con un nuevo nombre
documento_traducido.save('prueba_ingles.docx')
print("¡Éxito total! Se ha creado el archivo 'prueba_ingles.docx' en tu carpeta.")
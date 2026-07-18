import streamlit as st
import docx
from pypdf import PdfReader
import io
from deep_translator import GoogleTranslator
import time
import concurrent.futures

st.set_page_config(page_title="Traductor Profesional IA", page_icon="🌐")
st.title("🌐 Traductor de Documentos (Word & PDF)")

idiomas = {"Español": "es", "Inglés": "en", "Francés": "fr", "Alemán": "de", "Italiano": "it"}
col1, col2 = st.columns(2)
with col1: orig = st.selectbox("Idioma de origen", list(idiomas.keys()))
with col2: dest = st.selectbox("Idioma de destino", list(idiomas.keys()))

archivo = st.file_uploader("Sube tu archivo (Word o PDF)", type=['docx', 'pdf'])

def translate_batch(batch_text, source_lang='es', target_lang='en', max_retries=3):
    translator = GoogleTranslator(source=source_lang, target=target_lang)
    for attempt in range(max_retries):
        try:
            return translator.translate(batch_text)
        except Exception:
            if attempt == max_retries - 1:
                return None
            time.sleep(2 ** attempt)

def translate_paragraphs_parallel(paragraphs, source_lang, target_lang, progress_callback=None):
    batches = []
    current_batch = []
    current_length = 0
    
    for p in paragraphs:
        p_clean = p.strip()
        if not p_clean:
            continue
        if current_length + len(p_clean) + 1 > 4000:
            batches.append(current_batch)
            current_batch = [p_clean]
            current_length = len(p_clean)
        else:
            current_batch.append(p_clean)
            current_length += len(p_clean) + 1
    if current_batch:
        batches.append(current_batch)
        
    if not batches:
        return {}
        
    translated_results = [None] * len(batches)
    total_batches = len(batches)
    completed_batches = 0
    
    max_workers = min(10, total_batches)
    
    with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as executor:
        futures = {
            executor.submit(translate_batch, "\n".join(batch), source_lang, target_lang): idx
            for idx, batch in enumerate(batches)
        }
        for future in concurrent.futures.as_completed(futures):
            idx = futures[future]
            translated_results[idx] = future.result()
            completed_batches += 1
            if progress_callback:
                progress_callback(completed_batches / total_batches)
                
    final_paragraphs_map = {}
    paragraph_index = 0
    
    for idx, batch in enumerate(batches):
        translated_text = translated_results[idx]
        if translated_text:
            lines = [line.strip() for line in translated_text.split('\n')]
            if len(lines) == len(batch):
                for line in lines:
                    final_paragraphs_map[paragraph_index] = line
                    paragraph_index += 1
            else:
                sub_results = [None] * len(batch)
                sub_workers = min(5, len(batch))
                with concurrent.futures.ThreadPoolExecutor(max_workers=sub_workers) as sub_executor:
                    sub_futures = {
                        sub_executor.submit(translate_batch, item, source_lang, target_lang): s_idx
                        for s_idx, item in enumerate(batch)
                    }
                    for sub_future in concurrent.futures.as_completed(sub_futures):
                        s_idx = sub_futures[sub_future]
                        sub_results[s_idx] = sub_future.result()
                for s_idx, item in enumerate(batch):
                    final_paragraphs_map[paragraph_index] = sub_results[s_idx] if sub_results[s_idx] else item
                    paragraph_index += 1
        else:
            for item in batch:
                final_paragraphs_map[paragraph_index] = item
                paragraph_index += 1
                
    return final_paragraphs_map

if archivo is not None:
    if st.button("Traducir ahora"):
        progress_bar = st.progress(0.0)
        status_text = st.empty()
        
        def update_progress(percent):
            progress_bar.progress(percent)
            status_text.text(f"Progreso de traducción: {int(percent * 100)}%")

        try:
            if archivo.name.endswith('.docx'):
                # Lectura in-situ para preservar formato, estilos e imágenes del documento
                doc = docx.Document(archivo)
                
                # Coleccionamos todas las referencias a párrafos en el cuerpo principal y tablas
                paragraphs_to_translate = []
                for p in doc.paragraphs:
                    if p.text.strip():
                        paragraphs_to_translate.append(p)
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                if p.text.strip():
                                    paragraphs_to_translate.append(p)
                
                texts = [p.text for p in paragraphs_to_translate]
                
                if texts:
                    translated_map = translate_paragraphs_parallel(
                        texts, idiomas[orig], idiomas[dest], update_progress
                    )
                    for idx, p in enumerate(paragraphs_to_translate):
                        if idx in translated_map:
                            p.text = translated_map[idx]
                
                memoria = io.BytesIO()
                doc.save(memoria)
                memoria.seek(0)
                
            elif archivo.name.endswith('.pdf'):
                # Lectura del texto del PDF
                lector = PdfReader(archivo)
                lista_parrafos = []
                for pagina in lector.pages:
                    texto_p = pagina.extract_text()
                    if texto_p:
                        lista_parrafos.extend(texto_p.split('\n'))
                
                non_empty_indices = [i for i, p in enumerate(lista_parrafos) if p.strip()]
                non_empty_texts = [lista_parrafos[i] for i in non_empty_indices]
                
                if non_empty_texts:
                    translated_map = translate_paragraphs_parallel(
                        non_empty_texts, idiomas[orig], idiomas[dest], update_progress
                    )
                    
                    final_list = list(lista_parrafos)
                    for idx_in_non_empty, original_idx in enumerate(non_empty_indices):
                        if idx_in_non_empty in translated_map:
                            final_list[original_idx] = translated_map[idx_in_non_empty]
                else:
                    final_list = lista_parrafos
                
                # PDF se exporta a Word (formato original de la app)
                doc_final = docx.Document()
                for p in final_list:
                    doc_final.add_paragraph(p)
                
                memoria = io.BytesIO()
                doc_final.save(memoria)
                memoria.seek(0)
                
            st.success("¡Traducción completa!")
            st.download_button("⬇️ Descargar archivo traducido (.docx)", memoria, "Traduccion.docx")
        except Exception as e:
            st.error(f"Ocurrió un error durante la traducción: {e}")